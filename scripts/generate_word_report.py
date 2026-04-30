from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_table_to_doc(doc, title, headers, data):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def create_word_doc():
    doc = Document()

    # Título Principal
    title = doc.add_heading('Modelo Relacional Normalizado - SAGE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Este documento presenta el diseño técnico de la base de datos para el Sistema de Gestión Escolar (SAGE), estructurado bajo la Tercera Forma Normal (3NF) para asegurar integridad y escalabilidad.')

    # 1. Justificación 3NF
    doc.add_heading('1. Justificación de Normalización (3NF)', level=1)
    doc.add_paragraph('La base de datos ha sido normalizada para eliminar anomalías de inserción, actualización y borrado:')
    bullet_1 = doc.add_paragraph('Eliminación de redundancia al separar perfiles de usuario y roles.', style='List Bullet')
    bullet_2 = doc.add_paragraph('Desvinculación de calificaciones y asistencias de las clases directas, asociándolas a la entidad "Inscripción".', style='List Bullet')
    bullet_3 = doc.add_paragraph('Uso sistemático de Llaves Primarias (PK) y Foráneas (FK) para todas las relaciones.', style='List Bullet')

    # 2. Diccionario de Datos
    doc.add_heading('2. Estructura de Tablas', level=1)
    
    headers = ['Campo', 'Llave', 'Descripción']
    
    # Usuarios
    add_table_to_doc(doc, 'Tabla: usuarios', headers, [
        ('id', 'PK', 'ID único del sistema.'),
        ('username', 'U', 'Nombre de acceso.'),
        ('password', '-', 'Hash de seguridad.'),
        ('rol_id', 'FK', 'Vínculo con roles.'),
        ('nombre_completo', '-', 'Nombre completo.')
    ])

    # Alumnos
    add_table_to_doc(doc, 'Tabla: alumnos', headers, [
        ('id', 'PK', 'Identificador académico.'),
        ('usuario_id', 'FK', 'Enlace a credenciales.'),
        ('carrera_id', 'FK', 'Especialidad cursada.'),
        ('semestre', '-', 'Ciclo actual.')
    ])

    # Clases
    add_table_to_doc(doc, 'Tabla: clases', headers, [
        ('id', 'PK', 'ID de la materia compartida.'),
        ('nombre_clase', '-', 'Nombre de la asignatura.'),
        ('maestro_id', 'FK', 'Profesor asignado (User ID).'),
        ('periodo', '-', 'Ej. 2024-1.')
    ])

    # Inscripciones
    add_table_to_doc(doc, 'Tabla: inscripciones', headers, [
        ('id', 'PK', 'Vínculo Alumno-Clase.'),
        ('alumno_id', 'FK', 'ID del alumno.'),
        ('clase_id', 'FK', 'ID de la clase.'),
        ('fecha', '-', 'Fecha de alta.')
    ])

    # Calificaciones
    add_table_to_doc(doc, 'Tabla: calificaciones', headers, [
        ('id', 'PK', 'Identificador de nota.'),
        ('inscripcion_id', 'FK', 'Referencia a la inscripción.'),
        ('evaluacion', '-', 'Tipo (Parcial/Final).'),
        ('valor', '-', 'Puntuación decimal.')
    ])

    # 3. Diagrama
    doc.add_heading('3. Modelo Relacional Visual', level=1)
    
    img_path = r'C:\Users\jefte\.gemini\antigravity\brain\8cb6a16b-7057-4bd4-8555-d9ed232c3321\sage_relational_model_3nf_1774210666352.png'
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Error al insertar imagen: {str(e)}]')
            doc.add_paragraph('Por favor, inserte manualmente el diagrama adjunto en el reporte.')
    
    # Mensaje final
    doc.add_paragraph('\n---\nGenerado por Antigravity AI para el proyecto SAGE.')

    # Guardar
    output_path = r'C:\Users\jefte\OneDrive\Documentos\SAGE\Modelo_Relacional_SAGE.docx'
    doc.save(output_path)
    print(f"Documento generado exitosamente en: {output_path}")

if __name__ == "__main__":
    create_word_doc()
